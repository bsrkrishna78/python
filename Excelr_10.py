#!/usr/bin/env python
# coding: utf-8

# In[1]:


get_ipython().system('pip install PyPDF2')


# In[6]:


import PyPDF2
from PyPDF2 import PdfFileReader


# In[4]:


PyPDF2.__version__


# In[8]:


import PyPDF2, urllib , nltk
from io import BytesIO
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords


# In[10]:


wFile=urllib.request.urlopen('https://www.udri.org/pdf/02%20working%20paper%201.pdf')
pdfreader=PyPDF2.PdfReader(BytesIO(wFile.read()))


# In[11]:


pageObj = pdfreader.pages[2]
page2 = pageObj.extract_text()
punctuations = ['(', ')', ':', ',', '[', ']', ';', ',', '.']
tokens = word_tokenize(page2)
stop_words = stopwords.words('english')
keywords = [word for word in tokens if not word in stop_words and not word in punctuations]


# In[12]:


keywords


# In[14]:


name_list=list()
check=['Mr.', 'Mrs.', 'Ms.']
for idx, token in enumerate(tokens):
    if token.startswith(tuple(check)) and idx<(len(tokens)-1):
        name=token+tokens[idx+1]+''+ tokens[idx+2]
        name_list.append(name)
        
print(name_list)


# In[15]:


wFile.close()


# In[20]:


pdf=open(r"C:\Users\bsrkr\Downloads\02 working paper 1.pdf","rb")
pdf_reader=PyPDF2.PdfReader(pdf)
print("Number of pages:",len(pdf_reader.pages))
page=pdf_reader.pages[1]
print(page.extract_text())
pdf.close()


# In[21]:


get_ipython().system('pip install python-docx')


# In[23]:


import docx


# In[28]:


doc=open("Converted_Document.docx","rb")
document=docx.Document(doc)


# In[31]:


docu = ""
for para in document.paragraphs:
    docu += para.text
print(docu)


# In[32]:


for i in range(len(document.paragraphs)):
    print("The content of the paragraph "+ str(i)+" is :" + document.paragraphs[i].text+"\n")


# In[33]:


get_ipython().system('pip install bs4')


# In[34]:


import urllib.request as urllib2
from bs4 import BeautifulSoup


# In[35]:


response = urllib2.urlopen('https://en.wikipedia.org/wiki/Natural_language_processing')
html_doc = response.read()


# In[ ]:




